from docx import Document
from docx.shared import Cm
import os

PATH_template = os.getcwd() + r'/static/N/'


def sing_detail(paragraph, name):
    # 表格中要先水平垂直居中，然后再添加图片
    paragraph.text = paragraph.text.replace(name, "")

    sign_image_path, height = {
        "徐伟": (PATH_template + r"徐伟.png", 0.7),
        "王冬冬": (PATH_template + r"王0.7.png", 0.7),
        "陆声来": (PATH_template + r"周1.1.png", 0.7),
        "竺潇杭": (PATH_template + r"周1.1.png", 0.7),
        "丁佳铭": (PATH_template + r"丁.png", 0.7),
    }[name]
    run = paragraph.add_run()
    run.add_picture(sign_image_path, height=Cm(height), width=None)
    print("签名替换成功")


def add_syt(doc, src):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "这是示意图" in paragraph.text:
                        # 表格中要先水平垂直居中，然后再添加图片
                        paragraph.text = paragraph.text.replace("这是示意图", "")
                        run = paragraph.add_run()
                        run.add_picture(src, height=Cm(5.61), width=None)
    return doc
