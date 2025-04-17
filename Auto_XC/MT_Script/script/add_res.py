from docx import Document
from docxcompose.composer import Composer
from .check_and_change import *
from docx.shared import Pt
import numpy as np


def add_res(doc_main, second_page, b_no, group_main):
    no_next = b_no  # 续报编号
    bn = 1
    doc_second = Document(second_page)  # 打开第二页模板,创建第二页的 document 对象
    cp = Composer(doc_main)  # 创建 Composer 对象
    date_grouped = group_main.groupby('检测日期')
    for date2, group in date_grouped:
        date = str(group["检测日期"].iloc[0].strftime("%Y.%m.%d"))  # 检测日期
        for row_idx, (row_i, row) in enumerate(group.iterrows(), start=1):
            dict_word = [
                row["焊缝编号"],
                row["检测部位(m)"],
                row["检测总长度(m)"],
                row["缺陷编号"],
                row["缺陷性质"],
                convert_to_int(row["X(mm)"]),
                row["缺陷尺寸(mm)"],
                row["结论"],
            ]  # 每行数据
            # date = str(row["检测日期"].strftime("%Y.%m.%d"))  # 检测日期
            add_line(doc_second, str((row_idx-1) % 25 + 1), dict_word)  # 添加当前行数据到第二页文档的 row_idx % 25 + 1 行
            if row_idx % 25 == 0:  # 每 25 行换页，报告编号要加续n，最后加页
                once_change(doc_second, {"这是检测日期": date})
                once_change(doc_second, {"这是报告的编号": no_next})  # res 页每一页都要替换报告编号
                cp.append(doc_second)  # 将结果的第二页文档添加到第一页主文档
                doc_second = Document(second_page)  # 重新打开第二页模板（重置）
                no_next = b_no + "（续" + str(bn) + "）"  # 为下一页的续报编号做准备
                bn += 1
        if len(group) % 25 != 0:  # 最后一页的处理
            once_change(doc_second, {"这是检测日期": date})
            once_change(doc_second, {"这是报告的编号": no_next})  # res 页每一页都要替换报告编号
            add_line(doc_second, str(len(group) % 25 + 1), ["以下空白"])
            cp.append(doc_second)  # 将结果的第二页文档添加到第一页主文档
            doc_second = Document(second_page)  # 重新打开第二页模板（重置）
            no_next = b_no + "（续" + str(bn) + "）"  # 为下一页的续报编号做准备
            bn += 1
    return doc_main


# 给第二页的文档某行添加数据
def add_line(doc, content, word):
    for table in doc.tables:
        for row in reversed(table.rows):
            if content in [cell.text for cell in row.cells]:
                for i in range(len(word)):
                    print(row.cells[0].text)
                    cell = row.cells[i + 1]
                    # 清空单元格中的所有段落内容
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.clear()  # 清除段落中的内容
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = (
                        cell.paragraphs[0].alignment if cell.paragraphs else None
                    )
                    if str(word[i]) == "nan":
                        run = paragraph.add_run("/")
                    else:
                        run = paragraph.add_run(str(word[i]))
                    run.font.name = 'Times New Roman'  # 设置字体为 Times New Roman
                    run.font.size = Pt(9)  # 设置字体大小为文档默认字体大小
                break


def convert_to_int(word):
    try:
        return int(word)
    except ValueError:
        return word
