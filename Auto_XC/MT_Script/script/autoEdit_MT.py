
import shutil
from .add_res import *
from .add_sign import *
from .dnf import *
from datetime import timedelta
from docx import Document
from .replace_the_dict import *
from multiprocessing import Pool
from colorama import Fore, init
from openpyxl import Workbook
from openpyxl.utils import get_column_letter
PATH_result = os.getcwd() + r'/Auto_XC/MT_Script/report_result/'
PATH_TEMPLATE = os.getcwd() + r'/Auto_XC/MT_Script/report_template/'
init(autoreset=True)


def flatten_nested_list(nested_list):
    result = []
    for item in nested_list:
        if isinstance(item[0], list):  # 检查是否是嵌套列表
            result.extend(flatten_nested_list(item))
        else:
            result.append(item)
    return result


def process_report(params):
    no, ((component, weld, g), group), pt_doc, res_page, bk1, res_path = params
    log_info = []
    document_copy = Document(pt_doc)
    total_length = str(round(group['检测总长度(m)'].sum(), 3))  # 计算总长度
    final_date = group['检测日期'].max().strftime('%Y.%m.%d')
    b_no = bk1["报告编号"][0][:-3] + f"{no:03}"  # 报告编号
    qx_num = (group['结论'] == '不合格').sum()
    print(f'报告编号：{no:<10}' f'构件名称: {component:<10}'f'焊缝类型: {weld:<10}'f'len: {len(group):<10}'
          f'总长度: {total_length:<10}'f'组：{g:<10}'f'不合格数量：{qx_num:<10}')
    save_path = res_path + "/" + str(b_no[-6:]) + " " + component.replace("/", "-") + ".docx"  # 保存路径
    save_path_retest = res_path + "/" + str(b_no[-6:]) + " " + component.replace("/", "-") + "-1.docx"  # 保存路径
    replace_eve_dict = {
        "这是构件名称": component, "这是检测部位": weld, "这是检测数量": total_length, "报告的日期": final_date,  # 替换字典
        "检测的结论": "检测结论合格", "这是报告的编号": b_no
    }
    if qx_num != 0:
        replace_eve_dict["检测的结论"] = f"发现 {qx_num} 处超标缺欠，其余焊缝检测结论合格"  # 超标缺陷数量不为0，报告结论为不合格
        print(Fore.RED + b_no + ": 有" + str(qx_num) + "不合格缺陷，报告结论为不合格")
    check_and_change(document_copy, replace_eve_dict)  # 替换第一页的内容
    document = add_res(document_copy, res_page, b_no, group)  # 添加结果页后的 document 对象
    final_datetime = datetime.datetime.strptime(final_date, '%Y.%m.%d')
    need_replace_time = datetime.datetime.strptime("2024.10.17", '%Y.%m.%d')
    if final_datetime < need_replace_time:  # 日期小于2024.10.17，需要替换公司名称
        replace_header_text(document, "视正检测", "视正钢结构检测")
        print(f"{b_no} 日期小于2024.10.17，需要替换公司名称")
    log_info.append(list(map(str, [no, b_no, final_date, component, weld, total_length, "合格"])))
    if qx_num != 0:  # 有不合格缺陷
        document_retest = Document(pt_doc)  # 打开修改后的模板第一页
        unqualified_grouped = group[group['结论'] == '不合格']
        final_date_next = (datetime.datetime.strptime(final_date, '%Y.%m.%d') + timedelta(days=1)).strftime('%Y.%m.%d')
        no_next = b_no + "-1"
        doc_second = Document(res_page)
        all_long = 0
        for row_idx, (row_i, row) in enumerate(unqualified_grouped.iterrows(), start=1):
            qxx = row["X(mm)"]
            qxl = row["缺陷尺寸(mm)"]
            l_all = row["检测总长度(m)"] * 1000
            if 50 <= qxx <= l_all - qxl - 50:  # 两边都够
                bw = str(round((qxx - 50) / 1000, 3)) + "~" + str(round((qxx + qxl + 50) / 1000, 3))
                zc = str(round((qxl + 100) / 1000, 3))
            elif qxx < 50 and qxx <= l_all - qxl - 50:  # 左边不够
                bw = "0~" + str(round((qxx + qxl + 50) / 1000, 3))
                zc = str(round((qxl + 50 + qxx) / 1000, 3))
            elif qxx >= 50 and qxx > l_all - qxl - 50:  # 右边不够
                bw = str(round((qxx - 50) / 1000, 3)) + "~" + str(round(l_all / 1000, 3))
                zc = str(round((l_all - qxx + 50) / 1000, 3))
            else:
                bw = "0~" + str(round((l_all) / 1000, 3))
                zc = str(round((l_all - qxx + 50) / 1000, 3))
            all_long += float(zc)
            dict_word = [
                row["焊缝编号"], bw, zc,
                '/', 'A', '/', '/', 'R1合格',
            ]  # 每行数据
            add_line(doc_second, str((row_idx - 1) % 25 + 1), dict_word)
        add_line(doc_second, str(len(unqualified_grouped) % 25 + 1), ["以下空白"])
        once_change(doc_second, {"这是检测日期": final_date_next})
        once_change(doc_second, {"这是报告的编号": no_next})  # res 页每一页都要替换报告编号

        replace_replace = {
            "这是构件名称": component, "这是检测部位": weld,
            "这是检测数量": round(all_long, 3),
            "报告的日期": final_date_next, "检测的结论": "检测结论合格",
            "这是报告的编号": no_next,
            "100%": '/', "20%": '/', "25%": '/'
        }
        check_and_change(document_retest, replace_replace)  # 替换复探第一页的内容
        if final_datetime < need_replace_time:  # 日期小于2024.10.17，需要替换公司名称
            replace_header_text(document, "视正检测", "视正钢结构检测")
            print(f"{b_no} 日期小于2024.10.17，需要替换公司名称")
        cp = Composer(document_retest)
        # cp.append(document_retest)  # 合并报告

        cp.append(doc_second)  # 将结果的第二页文档添加到第一页主文档
        final_date_next = datetime.datetime.strptime(final_date_next, '%Y.%m.%d')
        if final_date_next < need_replace_time:  # 日期小于2024.10.17，需要替换公司名称
            replace_header_text(doc_second, "视正检测", "视正钢结构检测")
            print(f"{b_no} 日期小于2024.10.17，需要替换公司名称")
        # doc_third = Document(r"C:\Users\95609\Desktop\出报告\中环312\五联附件.docx")
        # cp.append(doc_third) bin # 附件
        # once_change(document_retest, {"这是报告的编号": b_no})
        log_info.append([no, no_next, final_date_next, component, weld, all_long, "R1合格"])
        document_retest.save(save_path_retest)  # 保存文件

    # doc_third = Document(r"C:\Users\95609\Desktop\出报告\中环312\五联附件.docx")
    # cp = Composer(document)
    # cp.append(doc_third)  # 附件
    # once_change(document, {"这是报告的编号": b_no})
    document.save(save_path)  # 保存文件
    return log_info


def auto_edit_mt(book_path, report_type):
    print(Fore.GREEN + "正在处理 " + book_path + " 文件...")
    path_all = os.path.dirname(book_path) + "/"
    pt_doc = path_all + "now.docx"
    bk1 = pd.read_excel(book_path, engine='openpyxl', sheet_name="磁粉总体信息", index_col=0,
                        header=None).T.reset_index(drop=True)  # 读取 excel 表, 获取总体信息
    bk2 = pd.read_excel(book_path, engine='openpyxl', sheet_name="磁粉原始记录")  # 读取 excel 表, 获取台账信息
    print(report_type)
    if report_type == "yz-63-dn":
        first_page = path_all + "A磁粉现场-扬州-63-东南.docx"  # 第一页模板路径  "A超声现场-扬州.docx"
        res_page = path_all + "A磁粉检测结果页-扬州.docx"  # 第二页模板路径
    elif report_type == "yz-63-rt":
        first_page = path_all + "A磁粉现场-扬州-63-润通.docx"  # 第一页模板路径  "A超声现场-扬州.docx"
        res_page = path_all + "A磁粉检测结果页-扬州.docx"  # 第二页模板路径
    elif report_type == "yz-65-dn":
        first_page = path_all + "A磁粉现场-扬州-65-东南.docx"  # 第一页模板路径  "A超声现场-扬州.docx"
        res_page = path_all + "A磁粉检测结果页-扬州.docx"  # 第二页模板路径
    elif report_type == "yz-65-wx":
        first_page = path_all + "A磁粉现场-扬州-65-纬信.docx"  # 第一页模板路径  "A超声现场-扬州.docx"
        res_page = path_all + "A磁粉检测结果页-扬州.docx"  # 第二页模板路径
    elif report_type == "sz":
        first_page = path_all + "A磁粉现场.docx"  # 第一页模板路径  "A超声现场-扬州.docx"
        res_page = path_all + "A磁粉检测结果页.docx"  # 第二页模板路径
    else:
        print(Fore.RED + f"{report_type}暂不支持！")
        return
    res_path = PATH_result + bk1["报告编号"][0] + "/"  # 删除已存在的非空结果文件夹
    if os.path.exists(res_path):
        shutil.rmtree(res_path)  # 删除整个结果文件夹
    os.makedirs(res_path)  # 创建结果文件夹

    document = Document(first_page)  # 打开第一页
    replace_dict = {}  # 定义一个空字典来存储替换内容
    replace_the_dict(bk1, replace_dict)  # 设置第一页的替换字典
    check_and_change(document, replace_dict)  # 替换第一页的内容
    p_res = PATH_TEMPLATE + "示意图.png"
    add_syt(document, p_res)
    document.save(pt_doc)
    print(f'第一页模板报告已生成: {pt_doc}')
    grouped = dnf(bk2)  # 按 构件名称、焊缝类型、长度每超150m 分组
    # 准备传递给进程池的参数
    params_list = []
    log_list = []
    start_no = int(bk1["报告编号"][0][-3:])  # 起始报告编号
    end_no = start_no + len(grouped) - 1  # 结束报告编号

    for no, ((component, weld, g), group) in enumerate(grouped, start_no):
        params_list.append((no, ((component, weld, g), group), pt_doc, res_page, bk1, res_path))
    # 使用 multiprocessing
    with Pool(processes=4) as pool:  # 这里可以根据 CPU 核心数调整进程数
        log_list += pool.map(process_report, params_list)
    # 创建工作簿和工作表
    wb = Workbook()
    ws = wb.active
    ws.title = "检测报告数据"
    headers = ["序号", "报告编号", "报告日期", "构件名称", "检测部位", "检测数量", "结论"]
    ws.append(headers)  # 写入表头
    log_list = flatten_nested_list(log_list)
    for row_data in log_list:  # 使用for循环写入数据
        ws.append(row_data)
    # 自动调整列宽（优化显示）
    for col_idx, _ in enumerate(headers, 1):
        col_letter = get_column_letter(col_idx)
        ws.column_dimensions[col_letter].width = 15
    filename = res_path + str(start_no) + "~" + str(end_no) + ".xlsx"
    wb.save(filename)
    shutil.copy2(str(book_path), str(res_path))
    print('所有报告生成完成！')
    # global qx_res
    return res_path
